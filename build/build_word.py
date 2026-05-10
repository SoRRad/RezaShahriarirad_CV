"""Generate Shahriarirad_Reza_CV.docx from CSV source data."""
from __future__ import annotations

import datetime as _dt
import pathlib
import re
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))
from utils import author_tag_counts, get_profile, load_all_data, parse_semicolon  # noqa: E402

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = pathlib.Path(__file__).parent.parent
OUT = ROOT / "Shahriarirad_Reza_CV.docx"

FONT = "Arial"
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x1A, 0x4D, 0x7C)
GOLD = RGBColor(0xB8, 0x95, 0x3A)
GRAY = RGBColor(0x5D, 0x68, 0x78)
LIGHT_BLUE = "D9EAF7"
LIGHT_GOLD = "F7F0DC"


def _clean(value) -> str:
    return " ".join(str(value or "").replace("\u00a0", " ").split())


def _set_run(run, size=9.5, bold=False, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _paragraph(cell_or_doc, text="", size=9.5, bold=False, italic=False, color=None, after=2, before=0):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    run = p.add_run(_clean(text))
    _set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _cell_shade(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _cell_text(cell, text, size=9, bold=False, italic=False, color=None, shade=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if shade:
        _cell_shade(cell, shade)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(_clean(text))
    _set_run(run, size=size, bold=bold, italic=italic, color=color)


def _table(doc, rows, headers=None, widths=None):
    n_cols = len(headers or rows[0])
    table = doc.add_table(rows=1 if headers else 0, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if headers:
        for idx, header in enumerate(headers):
            _cell_text(table.rows[0].cells[idx], header, size=8.5, bold=True, color=NAVY, shade=LIGHT_BLUE)
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            _cell_text(row.cells[idx], value, size=8.7)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    _set_run(run, size=10.5, bold=True, color=NAVY)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    border.append(bottom)
    p_pr.append(border)


def _subheading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    _set_run(run, size=9.5, bold=True, color=BLUE)
    return p


def _public_email_records(profile):
    records = []
    if _clean(profile.get("email_professional_public_visible", "yes")).lower() == "yes":
        records.append(("Mayo Email", profile.get("email_professional", "")))
    if _clean(profile.get("email_personal_public_visible", "yes")).lower() == "yes":
        records.append(("Personal Email", profile.get("email_personal", "")))
    out, seen = [], set()
    for label, email in records:
        email = _clean(email)
        if email and email.lower() not in seen:
            seen.add(email.lower())
            out.append((label, email))
    return out


def _profile_links(profile):
    labels = [
        ("ORCID", profile.get("orcid_url") or profile.get("orcid")),
        ("Google Scholar", profile.get("scholar_url")),
        ("PubMed", profile.get("pubmed_url")),
        ("Scopus", profile.get("scopus_url")),
        ("ResearchGate", profile.get("researchgate_url")),
        ("Web of Science", profile.get("wos_url")),
        ("LinkedIn", profile.get("linkedin_url")),
        ("GitHub", profile.get("github_url")),
    ]
    return [(label, _clean(url)) for label, url in labels if _clean(url)]


def _set_default_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor(0x22, 0x2A, 0x35)


def _header(doc, profile):
    name = profile.get("name", "Reza Shahriarirad, M.D.")
    if not name.endswith("M.D."):
        name = "Reza Shahriarirad, M.D."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(name)
    _set_run(run, size=18, bold=True, color=NAVY)

    subtitle = f"{profile.get('title','')} | {profile.get('institution','')} | {profile.get('city_state','')}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(_clean(subtitle))
    _set_run(run, size=9.5, color=GRAY)

    emails = " | ".join(email for _, email in _public_email_records(profile))
    links = " | ".join(f"{label}: {url}" for label, url in _profile_links(profile)[:5])
    for line in [emails, links]:
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        _set_run(run, size=8.2, color=BLUE)


def _personal_information(doc, profile):
    _section_heading(doc, "Personal Information")
    rows = [
        ("Surname", "Shahriarirad"),
        ("Given name", "Reza"),
        ("Current residence", f"{profile.get('city_state','')}, USA"),
        ("Date of writing CV", _dt.date.today().strftime("%B %d, %Y")),
        ("Verified reviews / peer reviews", profile.get("peer_reviews", "")),
        ("Total citations", profile.get("citations_cached", "")),
        ("H-index", profile.get("h_index_cached", "")),
        ("Research interests", profile.get("research_interests", "")),
    ]
    _table(doc, rows, widths=[1.85, 4.75])


def _timeline_table(doc, title, df, columns):
    _section_heading(doc, title)
    rows = []
    for _, row in df.iterrows():
        values = []
        for key, label in columns:
            if key == "_role_org":
                values.append(_clean(f"{row.get('role','')} | {row.get('org','')}"))
            else:
                values.append(row.get(key, ""))
        rows.append(values)
    _table(doc, rows, headers=[label for _, label in columns])


def _awards(doc, data):
    _timeline_table(doc, "Awards", data["awards"], [("year", "Year"), ("title", "Award"), ("org", "Organization")])


def _leadership(doc, data):
    _timeline_table(
        doc,
        "Leadership & Community Service",
        data["leadership"],
        [("period", "Period"), ("title", "Role"), ("org", "Organization"), ("desc", "Description")],
    )


def _skills(doc, data):
    _section_heading(doc, "Skills")
    computing = "; ".join(
        f"{_clean(r.get('name'))} ({_clean(r.get('level'))})" if _clean(r.get("level")) else _clean(r.get("name"))
        for _, r in data["skills_computing"].iterrows()
    )
    interpersonal = "; ".join(_clean(r.get("name")) for _, r in data["skills_interpersonal"].iterrows())
    _table(doc, [("Computing", computing), ("Interpersonal", interpersonal)], widths=[1.4, 5.2])


def _patents(doc, data):
    _timeline_table(doc, "Patents and Innovations", data["patents"], [("date", "Date"), ("title", "Title"), ("issuer", "Issuer"), ("number", "Number")])


def _publication_metrics(doc, profile, pubs_df):
    _section_heading(doc, "Publication Metrics")
    records = pubs_df.to_dict("records")
    counts = author_tag_counts(records)
    rows = [
        ("Total publications", str(len(records))),
        ("First author", str(counts["first"])),
        ("Co-first author", str(counts["co-first"])),
        ("Last / senior author", str(counts["last"])),
        ("Corresponding author", str(counts["corresponding"])),
        ("H-index", profile.get("h_index_cached", "")),
        ("Total citations", profile.get("citations_cached", "")),
    ]
    _table(doc, rows, widths=[2.2, 4.4])


def _editor_reviewer(doc, data, profile):
    _section_heading(doc, "Editor & Reviewer Experience")
    editorial_rows = [
        (r.get("role", ""), r.get("journal", ""), r.get("period", "")) for _, r in data["editorial"].iterrows()
    ]
    _table(doc, editorial_rows, headers=["Role", "Journal", "Period"], widths=[1.7, 3.2, 1.2])
    review_summary = (
        f"{profile.get('peer_reviews','')} verified peer reviews; "
        f"{profile.get('manuscripts_reviewed','')} manuscripts reviewed; "
        f"{profile.get('journals_reviewed','')} journals reviewed."
    )
    _paragraph(doc, review_summary, size=9)


def _presentation_summary(doc, pres_df):
    _section_heading(doc, "Presentations Summary")
    type_counts = pres_df["type"].str.lower().value_counts().to_dict()
    rows = [
        ("Total presentations", str(len(pres_df))),
        ("Posters", str(type_counts.get("poster", 0))),
        ("Oral presentations", str(type_counts.get("oral", 0))),
    ]
    _table(doc, rows, widths=[2.2, 4.4])


def _languages_hobbies(doc, profile, hobbies_df):
    _section_heading(doc, "Languages, Extracurricular Activities and Hobbies")
    hobbies = "; ".join(_clean(r.get("name")) for _, r in hobbies_df.iterrows())
    _table(doc, [("Languages", profile.get("languages", "")), ("Interests", hobbies)], widths=[1.4, 5.2])


def _references(doc, refs_df):
    _section_heading(doc, "References")
    rows = []
    for _, r in refs_df.iterrows():
        links = []
        for idx in (1, 2):
            label = _clean(r.get(f"link_label_{idx}", ""))
            url = _clean(r.get(f"link_url_{idx}", ""))
            if label and url:
                links.append(f"{label}: {url}")
        rows.append((r.get("name", ""), r.get("role", ""), r.get("inst", ""), "; ".join(links)))
    _table(doc, rows, headers=["Name", "Role", "Institution", "Public links"], widths=[1.55, 2.0, 2.15, 1.25])


def _bold_author_name(paragraph, authors):
    parts = re.split(r"(Shahriarirad R\*?)", _clean(authors))
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        _set_run(run, size=8.4, bold=bool(re.fullmatch(r"Shahriarirad R\*?", part)))


def _appendix_publications(doc, pubs_df):
    doc.add_section(WD_SECTION.NEW_PAGE)
    _section_heading(doc, "Appendix A: Full Publication List")
    groups = [
        ("original", "Original Articles"),
        ("review", "Reviews / Meta-analyses"),
        ("case", "Case Reports"),
        ("letter", "Letters / Editorials"),
    ]
    for key, label in groups:
        subset = pubs_df[pubs_df["type"].str.lower() == key]
        if subset.empty:
            continue
        _subheading(doc, label)
        for _, r in subset.iterrows():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(2)
            n = p.add_run(f"{r.get('n','')}. ")
            _set_run(n, size=8.4, bold=True, color=NAVY)
            _bold_author_name(p, r.get("authors", ""))
            title = p.add_run(f" {_clean(r.get('title',''))}. ")
            _set_run(title, size=8.4)
            journal = p.add_run(_clean(r.get("journal", "")))
            _set_run(journal, size=8.4, italic=True)
            year = p.add_run(f" ({_clean(r.get('year',''))})")
            _set_run(year, size=8.4, color=GRAY)


def _appendix_journals(doc, journals_df):
    doc.add_section(WD_SECTION.NEW_PAGE)
    _section_heading(doc, "Appendix B: Journals Reviewed")
    names = [_clean(r.get("name")) for _, r in journals_df.iterrows() if _clean(r.get("name"))]
    half = (len(names) + 1) // 2
    rows = []
    for idx in range(half):
        rows.append((names[idx], names[idx + half] if idx + half < len(names) else ""))
    _table(doc, rows, widths=[3.3, 3.3])


def _appendix_presentations(doc, pres_df):
    doc.add_section(WD_SECTION.NEW_PAGE)
    _section_heading(doc, "Appendix C: Presentations")
    rows = []
    for idx, (_, r) in enumerate(pres_df.iterrows(), start=1):
        rows.append(
            (
                str(idx),
                r.get("date", ""),
                r.get("type", "").capitalize(),
                r.get("title", ""),
                f"{_clean(r.get('venue',''))}; {_clean(r.get('location',''))}",
            )
        )
    _table(doc, rows, headers=["#", "Date", "Type", "Title", "Venue / location"], widths=[0.35, 0.8, 0.7, 3.1, 2.0])


def main():
    data = load_all_data()
    profile = get_profile(data)

    doc = Document()
    _set_default_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.62)
        section.bottom_margin = Inches(0.62)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    _header(doc, profile)
    _personal_information(doc, profile)
    _timeline_table(doc, "Education", data["education"], [("period", "Period"), ("degree", "Degree"), ("org", "Institution"), ("city", "Location")])
    _timeline_table(doc, "Working and Professional Experience", data["experience"], [("period", "Period"), ("_role_org", "Role / organization"), ("city", "Location"), ("desc", "Description")])
    _awards(doc, data)
    _leadership(doc, data)
    _skills(doc, data)
    _patents(doc, data)
    _publication_metrics(doc, profile, data["publications"])
    _editor_reviewer(doc, data, profile)
    _presentation_summary(doc, data["presentations"])
    _languages_hobbies(doc, profile, data["hobbies"])
    _references(doc, data["references"])
    _appendix_publications(doc, data["publications"])
    _appendix_journals(doc, data["journals"])
    _appendix_presentations(doc, data["presentations"])

    doc.save(OUT)
    print(f"  Generated {OUT.relative_to(ROOT)} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
