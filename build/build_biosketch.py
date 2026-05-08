"""
build_biosketch.py — generate a NIH Non-Fellowship Biographical Sketch (.docx).
Run manually: python build/build_biosketch.py [--project "Project Name"]
Output: Shahriarirad_NIH_Biosketch_General.docx (repo root)
"""
import pathlib, sys, argparse, io, re
sys.path.insert(0, str(pathlib.Path(__file__).parent))
from utils import load_all_data, get_profile, parse_semicolon

ROOT = pathlib.Path(__file__).parent.parent

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, size_pt=11, bold=False, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    set_font(run, size_pt=11, bold=True)
    p.paragraph_format.border_bottom = True
    # Add bottom border via XML
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def bold_name(text):
    """Return list of (run_text, bold) pairs with applicant name bolded."""
    parts = re.split(r'(Shahriarirad R\*?)', text)
    return [(p, bool(re.match(r'Shahriarirad R\*?', p))) for p in parts if p]


def add_formatted_authors(para, authors):
    for text, is_bold in bold_name(authors):
        r = para.add_run(text)
        r.bold = is_bold
        r.font.size = Pt(10)


def select_pubs(pubs_df, category_keys, max_count=4, prefer_featured=True):
    """Select up to max_count publications from given cat keys."""
    filtered = pubs_df[pubs_df['cat'].apply(
        lambda c: bool(set(parse_semicolon(c)) & set(category_keys))
    )].copy()
    if prefer_featured and 'featured' in filtered.columns:
        featured = filtered[filtered['featured'].str.lower() == 'yes']
        rest     = filtered[filtered['featured'].str.lower() != 'yes'].sort_values('year', ascending=False)
        pool = pd.concat([featured, rest])
    else:
        pool = filtered.sort_values('year', ascending=False)
    # Prefer first/corresponding author
    auth_tags = ['first', 'corresponding', 'last']
    def auth_score(row):
        tags = parse_semicolon(row.get('tags',''))
        for i, t in enumerate(auth_tags):
            if t in tags:
                return i
        return 99
    pool['_ascore'] = pool.apply(auth_score, axis=1)
    pool = pool.sort_values(['_ascore','year'], ascending=[True,False])
    return pool.head(max_count)


# ── Document builder ───────────────────────────────────────────────────────

def build_biosketch(project_name=None):
    import pandas as pd
    data    = load_all_data()
    profile = get_profile(data)
    pubs_df = data["publications"]
    exp_df  = data["experience"]
    edu_df  = data["education"]

    doc = Document()

    # Page margins (NIH: 0.5" all sides)
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ── HEADER ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OMB No. 0925-0001 and 0925-0002 (Rev. 10/2021 Approved Through 09/30/2024)")
    set_font(r, size_pt=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIOGRAPHICAL SKETCH")
    set_font(r, size_pt=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Provide the following information for the Senior/Key Personnel and Other Significant Contributors.")
    set_font(r, size_pt=9)
    r2 = p.add_run(" Follow this format for each person. DO NOT EXCEED FIVE PAGES.")
    set_font(r2, size_pt=9, bold=True)

    doc.add_paragraph()

    # ── NAME & POSITION TABLE ─────────────────────────────────────────────
    t = doc.add_table(rows=2, cols=3)
    t.style = 'Table Grid'
    cells = t.rows[0].cells
    cells[0].text = f"NAME: {profile.get('name','Reza Shahriarirad, M.D.')}"
    cells[1].text = ""
    cells[2].text = f"eRA COMMONS USER NAME: shahriarirad_r"
    cells = t.rows[1].cells
    cells[0].text = f"POSITION TITLE: {profile.get('title','Research Fellow')}"
    cells[1].text = ""
    cells[2].text = f"EDUCATION/TRAINING (Begin with baccalaureate or other initial professional education)"

    doc.add_paragraph()

    # ── EDUCATION/TRAINING TABLE ──────────────────────────────────────────
    add_section_heading(doc, "EDUCATION/TRAINING")
    edu_table = doc.add_table(rows=1, cols=4)
    edu_table.style = 'Table Grid'
    hdr = edu_table.rows[0].cells
    for i, h in enumerate(["INSTITUTION AND LOCATION", "DEGREE", "Completion Date MM/YYYY", "FIELD OF STUDY"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for _, r in edu_df.iterrows():
        row = edu_table.add_row().cells
        row[0].text = f"{r.get('org','')} — {r.get('city','')}"
        row[1].text = r.get('degree','')
        row[2].text = r.get('period','')
        row[3].text = r.get('field','Medicine')

    doc.add_paragraph()

    # ── SECTION A: PERSONAL STATEMENT ────────────────────────────────────
    add_section_heading(doc, "A. Personal Statement")

    if project_name:
        ps_intro = (
            f"I am applying for support for the project titled "{project_name}." "
            f"My research background and clinical training make me well-suited to contribute to this work."
        )
    else:
        ps_intro = (
            "As a physician-researcher and Research Fellow at the Surgery Innovation Center, Mayo Clinic, "
            "I bring a broad background in surgical outcomes research, minimally invasive surgery, "
            "and the application of artificial intelligence to clinical problems."
        )

    p = doc.add_paragraph(ps_intro)
    p.paragraph_format.space_after = Pt(6)

    bio2 = profile.get("bio_paragraph_2","")
    if bio2:
        # Strip HTML tags for docx
        bio2_clean = re.sub(r'<[^>]+>', '', bio2)
        doc.add_paragraph(bio2_clean)

    # 4 selected publications for Personal Statement
    all_cat_keys = [k for g in ['surgery','ai'] for k in ['plastic','thoracic','gi','oncology','machine_learning']]
    ps_pubs = pubs_df.copy()
    if 'featured' in ps_pubs.columns:
        ps_pubs_feat = ps_pubs[ps_pubs['featured'].str.lower()=='yes'].head(4)
        if len(ps_pubs_feat) < 4:
            rest = ps_pubs[ps_pubs['featured'].str.lower()!='yes'].sort_values('year',ascending=False)
            ps_pubs_sel = pd.concat([ps_pubs_feat, rest]).head(4)
        else:
            ps_pubs_sel = ps_pubs_feat.head(4)
    else:
        ps_pubs_sel = ps_pubs.sort_values('year', ascending=False).head(4)

    if len(ps_pubs_sel):
        p = doc.add_paragraph("Selected Publications:")
        p.runs[0].bold = True
        for _, pub in ps_pubs_sel.iterrows():
            p = doc.add_paragraph(style='List Number')
            add_formatted_authors(p, pub.get('authors',''))
            r = p.add_run(f". {pub.get('title','')}. ")
            r.font.size = Pt(10)
            r2 = p.add_run(pub.get('journal',''))
            r2.italic = True
            r2.font.size = Pt(10)
            year = pub.get('year','')
            url  = pub.get('url','')
            suffix = f". {year}."
            if url:
                suffix += f" {url}"
            r3 = p.add_run(suffix)
            r3.font.size = Pt(10)

    doc.add_paragraph()

    # ── SECTION B: POSITIONS AND HONORS ──────────────────────────────────
    add_section_heading(doc, "B. Positions, Scientific Appointments, and Honors")

    p = doc.add_paragraph("Positions and Scientific Appointments")
    p.runs[0].bold = True

    for _, r in exp_df.iterrows():
        p = doc.add_paragraph()
        r1 = p.add_run(f"{r.get('period','')}  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(f"{r.get('role','')}, {r.get('org','')}")
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # ── SECTION C: CONTRIBUTIONS TO SCIENCE ──────────────────────────────
    add_section_heading(doc, "C. Contribution to Science")

    themes = [
        {
            "title": "Surgical Outcomes and Innovation",
            "desc": (
                "My work in thoracic, vascular, and reconstructive surgery has generated a body of evidence "
                "guiding clinical practice and surgical technique. These studies span retrospective outcome analyses, "
                "randomised controlled trials, and systematic reviews in high-impact surgical journals."
            ),
            "cat_keys": ["thoracic","vascular","plastic","gi","endocrine","ortho"],
        },
        {
            "title": "Artificial Intelligence and Computational Medicine",
            "desc": (
                "I have applied machine learning and deep learning approaches to clinical prediction problems "
                "in surgical oncology and critical care, contributing to the growing intersection of AI and surgery."
            ),
            "cat_keys": ["machine_learning","computer_vision","oncology"],
        },
    ]

    import pandas as pd
    for i, theme in enumerate(themes, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {theme['title']}")
        r.bold = True
        r.font.size = Pt(10)
        doc.add_paragraph(theme["desc"])

        sel = select_pubs(pubs_df, theme["cat_keys"], max_count=4)
        sel_sorted = sel.sort_values('year', ascending=True)
        for _, pub in sel_sorted.iterrows():
            p = doc.add_paragraph(style='List Number')
            add_formatted_authors(p, pub.get('authors',''))
            r = p.add_run(f". {pub.get('title','')}. ")
            r.font.size = Pt(10)
            r2 = p.add_run(pub.get('journal',''))
            r2.italic = True
            r2.font.size = Pt(10)
            year = pub.get('year','')
            url  = pub.get('url','')
            suffix = f". {year}."
            if url:
                suffix += f" {url}"
            r3 = p.add_run(suffix)
            r3.font.size = Pt(10)
        doc.add_paragraph()

    # ── SECTION D: ADDITIONAL INFORMATION ────────────────────────────────
    add_section_heading(doc, "D. Additional Information: Research Support and/or Scholastic Performance")

    grants = profile.get("grants","").strip()
    p = doc.add_paragraph("Ongoing Research Support")
    p.runs[0].bold = True
    if grants:
        doc.add_paragraph(grants)
    else:
        doc.add_paragraph("None")

    p2 = doc.add_paragraph("Completed Research Support")
    p2.runs[0].bold = True
    doc.add_paragraph("None")

    # ── Save ──────────────────────────────────────────────────────────────
    out_path = ROOT / "Shahriarirad_NIH_Biosketch_General.docx"
    doc.save(str(out_path))
    print(f"  Saved NIH Biosketch: {out_path.name} ({out_path.stat().st_size:,} bytes)")


def main():
    parser = argparse.ArgumentParser(description="Generate NIH Non-Fellowship Biosketch.")
    parser.add_argument("--project", type=str, default=None,
                        help="Project name to reference in Personal Statement framing.")
    args = parser.parse_args()
    build_biosketch(project_name=args.project)


if __name__ == "__main__":
    main()
